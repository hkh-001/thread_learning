# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_heading_style(run, font_size=16, bold=True):
    """设置标题样式"""
    font = run.font
    font.size = Pt(font_size)
    font.bold = bold
    font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def set_body_style(run, font_size=12):
    """设置正文样式"""
    font = run.font
    font.size = Pt(font_size)
    font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_paragraph(doc, text, level=1):
    """添加标题段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_heading_style(run, font_size=16, bold=True)
    elif level == 2:
        set_heading_style(run, font_size=14, bold=True)
    else:
        set_heading_style(run, font_size=12, bold=True)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p

def add_body_paragraph(doc, text, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
    run = p.add_run(text)
    set_body_style(run, font_size=12)
    p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    return p

def add_code_block(doc, code_text):
    """添加代码块（使用等宽字体）"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Pt(24)
    return p

# 读取现有文档
doc = Document('中国矿业大学计算机学院实验报告(实验1模板) .docx')

# 清除现有内容（保留第一页可能有的表头信息）
# 如果需要替换所有内容，可以清空所有段落
# 这里我们添加到文档末尾

# 添加实验目的
add_heading_paragraph(doc, '一、实验目的', level=1)

purpose_text = """本实验旨在通过实现四种经典的图像卷积滤波器（Sobel、Scharr、Laplacian、Gaussian），深入理解 CPU 多核并行编程的基本原理与方法。具体目标包括：掌握 POSIX 线程库（pthread）的线程创建、同步与回收机制；理解数据划分与任务调度在并行计算中的重要性；分析图像卷积操作的并行化可行性，探索行级并行划分策略的实现细节；通过对比串行与并行执行效率，体会并行加速的实际效果；同时熟悉跨平台开发中常见的环境配置与兼容性问题处理方法。"""

add_body_paragraph(doc, purpose_text)

# 添加空行
doc.add_paragraph()

# 添加实验内容
add_heading_paragraph(doc, '二、实验内容', level=1)

# 1. 题目描述
add_heading_paragraph(doc, '1. 题目描述', level=2)

desc_text = """本次实验要求利用 pthread 库实现神经网络中常用卷积操作的并行加速。图像卷积是计算机视觉的基础操作，涉及大量矩阵乘法运算，计算密集度高，天然适合并行化。实验实现了四种不同的 3×3 卷积核：

Sobel 算子：用于边缘检测，通过计算 x、y 方向的梯度幅值来识别图像边缘；
Scharr 算子：Sobel 的改进版本，权重更大，对斜向边缘检测更精确；
Laplacian 算子：二阶导数算子，用于检测边缘和角点；
Gaussian 算子：平滑滤波器，用于图像去噪和模糊处理。

程序需支持 PNG、JPG、BMP 等常见图像格式的输入输出，自动转换为灰度图进行处理，并能够根据用户指定的线程数进行并行加速。"""

add_body_paragraph(doc, desc_text)

doc.add_paragraph()

# 2. 算法描述
add_heading_paragraph(doc, '2. 算法描述', level=2)

# 2.1 图像加载与预处理
add_heading_paragraph(doc, '（1）图像加载与预处理', level=3)

algo_text1 = """程序使用 stb_image 单头文件库实现图像加载，支持自动识别文件格式。加载后，若为彩色图像（RGB/RGBA），则按照标准灰度转换公式（Gray = 0.299R + 0.587G + 0.114B）转换为单通道灰度图。这一预处理步骤在单线程中完成，为后续并行卷积准备数据。"""

add_body_paragraph(doc, algo_text1)

# 2.2 卷积运算原理
add_heading_paragraph(doc, '（2）卷积运算原理', level=3)

algo_text2 = """卷积操作通过 3×3 滑动窗口在图像上移动，将窗口内像素与卷积核对应元素相乘后求和，得到输出图像对应位置的像素值。对于图像边界，采用零填充（Zero Padding）策略，即当卷积核超出图像边界时，越界位置的像素值视为 0。这种处理方式实现简单，且不会引入额外的边界复杂性。"""

add_body_paragraph(doc, algo_text2)

# 2.3 并行化策略
add_heading_paragraph(doc, '（3）并行化策略：按行划分', level=3)

algo_text3 = """本实验采用"按图像行划分"的任务分配策略。设图像高度为 H，线程数为 N，则将 H 行图像划分为 N 个连续的子区间。每个线程负责计算其分配行范围内的所有卷积操作。这种划分方式的优势在于：

数据访问局部性好：每个线程顺序访问连续的内存行，缓存命中率高；
无写冲突：各线程只写入输出图像的不同行区域，无需互斥锁等同步机制；
负载均衡实现简单：总行数除以线程数得到基本分配量，剩余行（H % N）依次分配给前几个线程，实现近似均衡。

具体实现中，ThreadArgs 结构体封装了每个线程所需的参数：输入图像指针、输出缓冲区指针、起始行号 start_row 和结束行号 end_row。主线程循环创建 pthread，传递对应参数，创建完成后调用 pthread_join 等待所有线程结束。"""

add_body_paragraph(doc, algo_text3)

# 2.4 四种卷积核实现
add_heading_paragraph(doc, '（4）四种卷积核的具体实现', level=3)

algo_text4 = """Sobel 算子包含两个卷积核，分别检测水平方向和垂直方向的边缘：

Gx = [[-1, 0, 1], [-2, 0, 2], [-1, 0, 1]]
Gy = [[-1, -2, -1], [0, 0, 0], [1, 2, 1]]

对每个像素，分别计算 Gx 和 Gy 方向的卷积结果，最终输出梯度幅值 sqrt(Gx² + Gy²)。由于梯度值范围可能超出 [0, 255]，需进行 min-max 归一化。

Scharr 算子与 Sobel 结构相同，但权重更大：

Gx = [[-3, 0, 3], [-10, 0, 10], [-3, 0, 3]]
Gy = [[-3, -10, -3], [0, 0, 0], [3, 10, 3]]

更大的权重使其对 45 度角的边缘响应更强，结果相对更平滑。计算流程与 Sobel 一致，同样需要归一化。

Laplacian 算子是二阶微分算子，使用单卷积核：

L = [[0, 1, 0], [1, -4, 1], [0, 1, 0]]

卷积结果取绝对值后，同样需要进行 min-max 归一化到 0-255 范围。

Gaussian 算子用于平滑，采用归一化的高斯核：

G = [[1/16, 2/16, 1/16], [2/16, 4/16, 2/16], [1/16, 2/16, 1/16]]

由于是平滑操作，输出值已在合理范围内，直接截断到 [0, 255] 即可，无需基于 min-max 的归一化。"""

add_body_paragraph(doc, algo_text4)

# 2.5 线程生命周期管理
add_heading_paragraph(doc, '（5）线程生命周期管理', level=3)

algo_text5 = """主函数中，首先根据用户指定的线程数创建 pthread_t 数组和 ThreadArgs 数组。通过计算确定每个线程处理的行范围：rows_per_thread = height / num_threads，余数 rem = height % num_threads。前 rem 个线程多处理一行，确保负载均衡。使用 pthread_create 创建线程，线程函数为 xxxWorker（如 sobelWorker），参数为 ThreadArgs 指针。线程创建完成后，主线程调用 pthread_join 循环等待所有工作线程结束，最后进行结果归一化并保存图像。"""

add_body_paragraph(doc, algo_text5)

doc.add_paragraph()

# 3. 运行截图
add_heading_paragraph(doc, '3. 运行截图', level=2)

screenshot_text = """【此处插入程序命令行运行截图：显示在终端中执行 Gaussian.exe t53.png output.png 4 的命令及输出结果，包括"Gaussian 并行卷积完成""线程数: 4""耗时: X ms"等提示信息】

【此处插入输入输出图像对比图：左侧为原始测试图像 t53.png，右侧为经 Sobel 处理后的 t53_sobel.png，展示边缘检测效果】

【此处插入四个算子处理结果对比图：在同一行展示原图、Sobel 边缘检测、Scharr 边缘检测、Laplacian 边缘检测、Gaussian 模糊的效果对比】

【此处插入不同线程数耗时对比表或柱状图：展示使用 1 线程、2 线程、4 线程、8 线程处理同一图像时的耗时变化趋势，体现并行加速效果】"""

add_body_paragraph(doc, screenshot_text, indent=False)

doc.add_paragraph()

# 4. 调试情况
add_heading_paragraph(doc, '4. 调试情况', level=2)

debug_text = """在实验开发过程中，遇到了以下几类典型问题：

（1）pthread 库在 Windows 环境下的配置问题。由于实验环境为 Windows 系统，而 pthread 是 POSIX 标准库，需要使用 MinGW 等支持 pthread 的编译器。在编译时需要添加 -pthread 链接选项，否则会出现 pthread_create 未定义引用的链接错误。最终通过 g++ -o Gaussian.exe Gaussian.cpp -pthread -std=c++11 命令成功编译。

（2）C++ 标准兼容性问题。原始代码中使用了 std::clamp（C++17 特性）和结构化绑定 auto [min_it, max_it] = ...（C++17），但实验环境的 GCC 版本较低（4.9.2），仅支持 C++11。因此需要将 std::clamp 替换为三元运算符表达式 v < 0 ? 0 : (v > 255 ? 255 : v)，将结构化绑定替换为传统的 std::minmax_element 返回 pair 对象的访问方式。

（3）图像路径与编码问题。Windows 命令行对中文路径支持不佳，在测试时需要将图像文件放置在英文路径下，或使用引号包裹完整路径。此外，原始代码仅支持 PGM 格式，为了方便测试和结果查看，使用 stb_image 库扩展支持了 PNG、JPG、BMP 格式，实现了自动格式识别和灰度转换。

（4）多线程任务划分边界问题。初期实现时未正确处理图像行数不能被线程数整除的情况，导致部分行未被处理或数组越界。修复方法是计算余数 rem = height % num_threads，前 rem 个线程多分配一行（extra = 1），确保所有行都被覆盖且负载相对均衡。

（5）归一化策略的选择问题。初期统一使用 min-max 归一化，但发现 Gaussian 模糊后的图像对比度被拉伸，不符合模糊预期。经分析，Gaussian 卷积核权重和为 1，输出值范围与输入相近，应直接截断而非归一化；而 Sobel、Scharr、Laplacian 的梯度值范围不确定，必须使用 min-max 归一化才能正确显示。因此针对不同算子采用了不同的后处理策略。"""

add_body_paragraph(doc, debug_text)

doc.add_paragraph()
doc.add_paragraph()

# 添加实验体会
add_heading_paragraph(doc, '三、实验体会', level=1)

conclusion_text = """通过本次实验，我对 CPU 多核并行编程有了实质性的理解。在编写 pthread 代码的过程中，我深刻体会到"数据划分"是并行程序设计的核心——对于图像卷积这类"每个输出像素独立计算"的任务，按行划分是一种简单有效的策略，既保证了负载均衡，又避免了复杂的线程同步问题。

实验中遇到的环境配置和兼容性问题让我认识到，实际开发不能仅仅关注算法本身，还需要考虑编译器版本、操作系统差异、第三方库的集成等工程问题。特别是将 C++17 代码降级到 C++11 的过程中，我意识到使用现代语言特性虽然方便，但在受限环境中保持代码的可移植性同样重要。

对比四种卷积核的实现，我理解了不同滤波器的设计思想：Sobel 和 Scharr 通过计算一阶梯度检测边缘，Laplacian 通过二阶导数检测零交叉点，Gaussian 则通过加权平均实现平滑。它们虽然都是卷积操作，但数学原理和结果处理策略各不相同。这种"同一框架、不同内核"的设计模式在实际工程中非常常见。

最后，通过测试不同线程数下的执行时间，我观察到并行加速比并非线性增长。当线程数超过 CPU 物理核心数后，由于上下文切换开销，加速效果会趋于平缓甚至下降。这让我理解了 Amdahl 定律的实际意义：程序的加速比受限于串行部分的比例，盲目增加线程数并不能无限提升性能。

本次实验不仅巩固了 pthread 编程的实践经验，更重要的是建立了"识别可并行任务—设计数据划分策略—实现线程级并行—分析性能瓶颈"的完整思维模式，为后续学习 GPU 并行计算（CUDA）打下了坚实基础。"""

add_body_paragraph(doc, conclusion_text)

# 保存文档
output_filename = '实验报告_并行卷积实现.docx'
doc.save(output_filename)
print(f"实验报告已保存至: {output_filename}")
